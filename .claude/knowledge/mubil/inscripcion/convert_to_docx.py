import os
import re
import sys
import subprocess

# Auto-instalar python-docx si no está presente
try:
    import docx
except ImportError:
    print("Instalando python-docx para la conversión...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_markdown_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: No existe el archivo {md_path}")
        return

    doc = Document()

    # Ajustar márgenes estándar de 2.54 cm (1 pulgada)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Estilos básicos
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Gris oscuro para lectura premium

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_rows = []
    
    # Expresión regular para enlaces de Markdown
    link_re = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')

    # Procesar línea a línea
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # Ignorar bloques de diagrama Mermaid
        if line.startswith("```mermaid"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            i += 1
            continue

        # Ignorar otros bloques de código
        if line.startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            i += 1
            continue

        # Separador horizontal
        if line == "---":
            doc.add_paragraph().paragraph_format.space_before = Pt(6)
            i += 1
            continue

        # Encabezados
        if line.startswith("# "):
            text = line[2:].strip()
            p = doc.add_heading(text, level=1)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            # Cambiar color del encabezado a azul oscuro
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
            i += 1
            continue
        elif line.startswith("## "):
            text = line[3:].strip()
            p = doc.add_heading(text, level=2)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.color.rgb = RGBColor(0x00, 0x55, 0x80)
            i += 1
            continue
        elif line.startswith("### "):
            text = line[4:].strip()
            p = doc.add_heading(text, level=3)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.color.rgb = RGBColor(0x33, 0x66, 0x99)
            i += 1
            continue

        # Detección y procesamiento de tablas
        if line.startswith("|"):
            in_table = True
            table_rows.append(line)
            i += 1
            continue
        elif in_table and not line.startswith("|"):
            # Fin de la tabla, renderizarla en Word
            in_table = False
            render_docx_table(doc, table_rows)
            table_rows = []
            # Continuar sin avanzar i, para procesar la línea actual

        # Elementos de lista
        if line.startswith("* ") or line.startswith("- "):
            text = line[2:].strip()
            # Limpiar negritas de Markdown (**texto**)
            text_clean = text.replace("**", "")
            # Limpiar enlaces
            text_clean = link_re.sub(r'\1 (\2)', text_clean)
            p = doc.add_paragraph(text_clean, style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        if line == "":
            i += 1
            continue

        # Párrafo normal
        # Limpiar negritas y enlaces simples
        text_clean = line.replace("**", "")
        text_clean = link_re.sub(r'\1', text_clean)
        p = doc.add_paragraph(text_clean)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        i += 1

    # Guardar
    doc.save(docx_path)
    print(f"Creado con éxito: {docx_path}")

def render_docx_table(doc, raw_rows):
    # Parsear celdas de las filas
    rows_data = []
    for r in raw_rows:
        # Dividir por '|' y limpiar espacios
        cols = [c.strip() for c in r.split("|")[1:-1]]
        rows_data.append(cols)

    if not rows_data:
        return

    # Filtrar la fila de separación (--- | ---)
    final_rows = []
    for row in rows_data:
        # Si tiene guiones en todas las celdas, es la fila separadora de cabecera
        if all(re.match(r'^[-:\s]+$', cell) for cell in row):
            continue
        final_rows.append(row)

    if not final_rows:
        return

    num_cols = len(final_rows[0])
    num_rows = len(final_rows)

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'

    # Rellenar
    for r_idx, row in enumerate(final_rows):
        for c_idx, cell_value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            # Limpiar negritas de la celda
            clean_value = cell_value.replace("**", "")
            cell.text = clean_value
            # Si es cabecera, poner negrita y fondo gris (si el word lo permite)
            if r_idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    # Añadir un espacio después de la tabla
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)

if __name__ == "__main__":
    base_dir = r"c:\Users\imanol\projects\imanol\saas\maps\.claude\knowledge\mubil\inscripcion"
    
    # 1. Resumen Ejecutivo
    md_resumen = os.path.join(base_dir, "Resumen_ejecutivo_eStrata_MUBIL_2026_v2_PROPUESTA.md")
    docx_resumen = os.path.join(base_dir, "Resumen_ejecutivo_eStrata_MUBIL_2026_v2_PROPUESTA.docx")
    print(f"Procesando resumen: {md_resumen}...")
    parse_markdown_to_docx(md_resumen, docx_resumen)

    # 2. Memoria Técnica
    md_memoria = os.path.join(base_dir, "Memoria_ES_2026_V3_PROPUESTA.md")
    docx_memoria = os.path.join(base_dir, "Memoria_ES_2026_V3_PROPUESTA.docx")
    print(f"Procesando memoria: {md_memoria}...")
    parse_markdown_to_docx(md_memoria, docx_memoria)
